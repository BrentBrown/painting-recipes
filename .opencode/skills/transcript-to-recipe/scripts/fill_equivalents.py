#!/usr/bin/env python3
"""
fill_equivalents.py

Fills paint equivalents in a painting recipe markdown file and generates a
print-ready DOCX. Looks up each Source Paint by brand using paints.json,
strips common paint name suffixes before matching, and writes the results
back to the markdown file in-place.

Single file mode:
    python fill_equivalents.py <recipe.md>
    python fill_equivalents.py <recipe.md> --force

Folder mode (process all .md files in a directory):
    python fill_equivalents.py /path/to/recipes/
    python fill_equivalents.py /path/to/recipes/ --force

Batch mode (legacy flags):
    python fill_equivalents.py --recipes-dir ./Recipes --printables-dir ./Printables
    python fill_equivalents.py --recipes-dir ./Recipes --force
"""

import argparse
import json
import re
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SUFFIX_PATTERN = re.compile(
    r"\s+(Gloss|Contrast|Shade|Wash|Matte|Matt|Layer|Base|Dry|Technical)$",
    re.IGNORECASE,
)

NO_EQ = "No equivalent"

# ---------------------------------------------------------------------------
# Paint name normalisation
# ---------------------------------------------------------------------------


def strip_suffixes(name: str) -> str:
    """Repeatedly strip trailing paint-name suffixes until stable."""
    stripped = name.strip()
    while True:
        new = SUFFIX_PATTERN.sub("", stripped)
        if new == stripped:
            break
        stripped = new
    return stripped


def strip_parenthetical(name: str) -> str:
    """Strip trailing parenthetical notes like (contrast), (spray), (undiluted)."""
    return re.sub(r"\s*\([^)]+\)\s*$", "", name).strip()


def infer_brand_from_paint(source_paint: str) -> str | None:
    """Infer brand for tables without a Brand column based on paint name hints."""
    if "(contrast)" in source_paint.lower():
        return "Citadel"
    # Wyldwood is a Citadel Contrast paint
    if source_paint.strip().lower() == "wyldwood":
        return "Citadel"
    return None


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------


def load_paints(script_dir: Path) -> dict:
    paints_path = script_dir / "paints.json"
    if not paints_path.exists():
        print(
            f"Error: paints.json not found at {paints_path}\n"
            f"Expected it alongside this script in: {script_dir}",
            file=sys.stderr,
        )
        sys.exit(1)
    try:
        with open(paints_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        print(f"Error: paints.json is malformed: {e}", file=sys.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# Spray primer lookup
# ---------------------------------------------------------------------------

SPRAY_PATTERN = re.compile(r"\b(spray|undercoat|primer)\b", re.IGNORECASE)


def build_spray_primer_lookup(paints: dict) -> dict:
    """Build reverse map: Citadel paint name -> Spray Primer name."""
    spray_table = paints.get("Spray Primers", {})
    reverse: dict[str, str] = {}
    for primer_name, entry in spray_table.items():
        cit_list = entry.get("citadel", [])
        if isinstance(cit_list, list):
            for cit_name in cit_list:
                reverse[cit_name] = primer_name
        elif isinstance(cit_list, str):
            reverse[cit_list] = primer_name
    return reverse


def is_spray_paint(source_paint: str) -> bool:
    """Check if the paint name indicates a spray/undercoat/primer."""
    return bool(SPRAY_PATTERN.search(source_paint))


def strip_spray_suffix(name: str) -> str:
    """Remove spray/undercoat/primer suffixes and parenthetical notes."""
    name = re.sub(r"\s*\(([^)]+)\)\s*$", "", name).strip()
    name = re.sub(r"\s+(spray|undercoat|primer)\s*$", "", name, flags=re.IGNORECASE).strip()
    return name


# ---------------------------------------------------------------------------
# Lookup logic
# ---------------------------------------------------------------------------


def build_wf_reverse_lookup(paints: dict) -> dict:
    """
    Build a reverse map from Citadel paint name → Warpaints Fanatic paint name.

    Where multiple WF paints share the same Citadel equivalent the names are
    joined with ' / '.  The 'No equivalent' key is ignored.
    """
    wf_table = paints.get("Warpaints Fanatic", {})
    reverse: dict[str, list[str]] = {}
    for wf_name, entry in wf_table.items():
        cit = entry.get("citadel")
        if cit and cit != NO_EQ:
            reverse.setdefault(cit, []).append(wf_name)
    return {cit: " / ".join(names) for cit, names in reverse.items()}


def lookup_equivalents(
    brand: str, source_paint: str, paints: dict, wf_reverse: dict, spray_primer_lookup: dict
) -> tuple:
    """
    Return (ttc_col, citadel_col, warpaints_fanatic_col) for a source paint.

    For brand "Two Thin Coats" the TTC column contains the source paint name
    (with wave); for all other brands the TTC column contains the looked-up
    TTC equivalent.

    The Warpaints Fanatic column is resolved via a reverse lookup from the
    Citadel equivalent; for a Warpaints Fanatic source paint the column
    contains the source paint itself.

    For Citadel Contrast and Army Painter Speedpaint, the Warpaints Fanatic
    column shows the cross-brand equivalent (Speedpaint or Citadel Contrast).

    When brand is "Citadel", also checks "Citadel Contrast" for Contrast paint
    mappings.
    """
    stripped = strip_suffixes(source_paint)

    brand_table = paints.get(brand)

    # --- Special handling: Check Citadel Contrast when brand is Citadel ---
    if brand == "Citadel":
        contrast_table = paints.get("Citadel Contrast")
        if contrast_table:
            contrast_entry = contrast_table.get(stripped) or contrast_table.get(
                source_paint.strip()
            )
            if contrast_entry:
                speedpaint = contrast_entry.get("speedpaint", NO_EQ)
                return NO_EQ, NO_EQ, speedpaint
        # Also check Citadel table for speedpaint field (e.g. Lahmian Medium)
        if brand_table:
            citadel_entry = brand_table.get(stripped) or brand_table.get(source_paint.strip())
            if citadel_entry:
                speedpaint = citadel_entry.get("speedpaint")
                if speedpaint:
                    return NO_EQ, source_paint.strip(), speedpaint

    # --- Special handling: Spray primers ---
    if is_spray_paint(source_paint):
        paint_base = strip_spray_suffix(source_paint)
        primer = spray_primer_lookup.get(paint_base)
        if primer:
            # TTC gets primer, Citadel keeps original source, WF gets primer
            return primer, source_paint.strip(), primer

    if brand_table is None:
        if brand not in ("Citadel Contrast", "Army Painter Speedpaint"):
            print(
                f"  Warning: Unknown brand '{brand}' — writing '{NO_EQ}' for all columns.",
                file=sys.stderr,
            )
        return NO_EQ, NO_EQ, NO_EQ

    # Try stripped name first, then original as fallback
    entry = brand_table.get(stripped) or brand_table.get(source_paint.strip())

    if entry is None:
        return NO_EQ, NO_EQ, NO_EQ

    # --- Special handling for Citadel Contrast ---
    if brand == "Citadel Contrast":
        speedpaint = entry.get("speedpaint", NO_EQ)
        return NO_EQ, NO_EQ, speedpaint

    # --- Special handling for Army Painter Speedpaint ---
    if brand == "Army Painter Speedpaint":
        contrast = entry.get("citadel_contrast", NO_EQ)
        return NO_EQ, NO_EQ, contrast

    # --- TTC column ---
    if brand == "Two Thin Coats":
        wave = entry.get("wave", "")
        ttc_col = f"{stripped} ({wave})" if wave else stripped
    else:
        ttc_name = entry.get("ttc") or NO_EQ
        if ttc_name == NO_EQ:
            ttc_col = NO_EQ
        else:
            wave = entry.get("wave", "")
            ttc_col = f"{ttc_name} ({wave})" if wave else ttc_name

    # --- Citadel column ---
    if brand == "Citadel":
        citadel_col = source_paint.strip()
    else:
        citadel_col = entry.get("citadel") or NO_EQ

    # --- Warpaints Fanatic column ---
    if brand == "Warpaints Fanatic":
        wf_col = source_paint.strip()
    else:
        wf_col = wf_reverse.get(citadel_col, NO_EQ)

    return ttc_col, citadel_col, wf_col


# ---------------------------------------------------------------------------
# Markdown table helpers
# ---------------------------------------------------------------------------


def parse_table_row(line: str) -> list | None:
    """Parse a markdown table row into a list of stripped cell strings."""
    stripped = line.strip()
    if not stripped.startswith("|"):
        return None
    cells = [c.strip() for c in stripped.split("|")]
    # Remove empty strings from leading/trailing '|'
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def is_separator_row(cells: list) -> bool:
    """Return True if every cell looks like a markdown separator (----)."""
    return bool(cells) and all(re.match(r"^:?-+:?$", c) for c in cells)


def format_table_row(cells: list) -> str:
    return "| " + " | ".join(cells) + " |\n"


# ---------------------------------------------------------------------------
# Speedpaint detection
# ---------------------------------------------------------------------------

SPEEDPAINT_NAMES: set | None = None


def get_speedpaint_names(paints: dict) -> set:
    """Get set of all Speedpaint paint names for detection."""
    global SPEEDPAINT_NAMES
    if SPEEDPAINT_NAMES is None:
        SPEEDPAINT_NAMES = set(paints.get("Army Painter Speedpaint", {}).keys())
    return SPEEDPAINT_NAMES


def is_speedpaint(paint_name: str, speedpaint_names: set) -> bool:
    """Check if a paint name (without suffix) is a Speedpaint."""
    if not paint_name or paint_name == NO_EQ:
        return False
    # Strip (SP) or (F) annotation before checking
    name = re.sub(r"\s*\((SP|F)\)$", "", paint_name).strip()
    return name in speedpaint_names


# ---------------------------------------------------------------------------
# Core fill logic
# ---------------------------------------------------------------------------


def fill_equivalents(lines: list, paints: dict, force: bool) -> tuple:
    """
    Process markdown lines, filling blank (or all-"No equivalent") paint
    equivalent cells using paints.json lookups.

    Supports 4, 5, and 6-column Paint Equivalents tables:
      6 col: Role | Brand | Source Paint | TTC | Citadel | WF/Speedpaint
      5 col: Role | Source Paint | TTC | Citadel | WF/Speedpaint
      4 col: Role | Source Paint | TTC | WF/Speedpaint

    For 5 and 4-column tables without a Brand column, the brand is inferred
    from paint name hints (e.g. "(contrast)" suffix → Citadel).

    Returns (updated_lines, rows_filled, rows_skipped).
    Idempotent: rows with all equivalent columns already populated are skipped
    unless --force is given.
    """
    speedpaint_names = get_speedpaint_names(paints)
    wf_reverse = build_wf_reverse_lookup(paints)
    spray_primer_lookup = build_spray_primer_lookup(paints)

    result = []
    in_equiv_section = False
    header_seen = False
    rows_filled = 0
    rows_skipped = 0
    warned_no_table = True
    header_line_idx = -1
    table_columns = 0

    for i, line in enumerate(lines):
        if re.match(r"^##\s+Paint Equivalents", line.strip()):
            in_equiv_section = True
            header_seen = False
            warned_no_table = False
            table_columns = 0
            result.append(line)
            continue

        if in_equiv_section and re.match(r"^##\s+", line.strip()):
            in_equiv_section = False

        if not in_equiv_section:
            result.append(line)
            continue

        cells = parse_table_row(line)

        if cells is None:
            result.append(line)
            continue

        if is_separator_row(cells):
            result.append(line)
            continue

        # Header row
        if not header_seen:
            if len(cells) >= 4 and cells[0].lower() == "role":
                header_seen = True
                table_columns = len(cells)
                header_line_idx = len(result)
            result.append(line)
            continue

        if len(cells) != table_columns:
            print(
                f"  Warning: Line {i + 1} has {len(cells)} column(s) (expected {table_columns}), "
                f"skipping: {line.rstrip()}",
                file=sys.stderr,
            )
            result.append(line)
            rows_skipped += 1
            continue

        # Unpack based on column count
        if table_columns == 6:
            role, brand, source_paint, ttc_val, citadel_val, wf_val = cells
        elif table_columns == 5:
            role, source_paint, ttc_val, citadel_val, wf_val = cells
            brand = ""
        elif table_columns == 4:
            role, source_paint, ttc_val, wf_val = cells
            brand = ""
            citadel_val = ""
        else:
            result.append(line)
            rows_skipped += 1
            continue

        # For brandless tables, infer brand from paint name hints
        if not brand and table_columns in (4, 5):
            clean_paint = strip_parenthetical(source_paint)
            inferred = infer_brand_from_paint(source_paint)
            if inferred:
                brand = inferred
                source_paint = clean_paint

        if not source_paint or not brand:
            result.append(line)
            continue

        all_filled = bool(ttc_val and wf_val and (table_columns < 5 or citadel_val))

        if all_filled and not force:
            rows_skipped += 1
            result.append(line)
            continue

        ttc_new, citadel_new, wf_new = lookup_equivalents(
            brand, source_paint, paints, wf_reverse, spray_primer_lookup
        )

        if table_columns == 6:
            new_cells = [role, brand, source_paint, ttc_new, citadel_new, wf_new]
        elif table_columns == 5:
            new_cells = [role, source_paint, ttc_new, citadel_new, wf_new]
        elif table_columns == 4:
            new_cells = [role, source_paint, ttc_new, wf_new]
        else:
            result.append(line)
            continue

        result.append(format_table_row(new_cells))
        rows_filled += 1

    # --- Determine WF column index ---
    if table_columns == 6:
        wf_col_idx: int | None = 5
    elif table_columns == 5:
        wf_col_idx = 4
    elif table_columns == 4:
        wf_col_idx = 3
    else:
        wf_col_idx = None

    # --- Scan result to classify WF column values ---
    has_fanatic = False
    has_speedpaint = False

    if wf_col_idx is not None:
        for line in result:
            cells = parse_table_row(line)
            if not cells or is_separator_row(cells):
                continue
            if len(cells) >= table_columns and cells[0].lower() != "role":
                raw = cells[wf_col_idx].strip() if cells[wf_col_idx] else ""
                # Strip any existing annotation before classifying
                wf_val = re.sub(r"\s*\((SP|F)\)$", "", raw).strip()
                if is_speedpaint(wf_val, speedpaint_names):
                    has_speedpaint = True
                elif wf_val and wf_val != NO_EQ:
                    has_fanatic = True

    # --- Update header label ---
    if has_speedpaint and has_fanatic:
        new_header = "Warpaints Fanatic / Speedpaint 2.0"
    elif has_speedpaint:
        new_header = "Speedpaint 2.0"
    else:
        new_header = None  # keep original

    if new_header and header_line_idx >= 0 and wf_col_idx is not None:
        hdr_cells = parse_table_row(result[header_line_idx])
        if hdr_cells and len(hdr_cells) > wf_col_idx:
            hdr_cells[wf_col_idx] = new_header
            result[header_line_idx] = format_table_row(hdr_cells)

    # --- Annotate with (F) / (SP) when both types are present ---
    if has_fanatic and has_speedpaint and wf_col_idx is not None:
        for i, line in enumerate(result):
            cells = parse_table_row(line)
            if not cells or is_separator_row(cells):
                continue
            if len(cells) >= table_columns and cells[0].lower() != "role":
                raw = cells[wf_col_idx].strip() if cells[wf_col_idx] else ""
                # Strip existing annotation before re-annotating
                wf_val = re.sub(r"\s*\((SP|F)\)$", "", raw).strip()
                if wf_val and wf_val != NO_EQ:
                    suffix = (
                        " (SP)" if is_speedpaint(wf_val, speedpaint_names) else " (F)"
                    )
                    cells[wf_col_idx] = wf_val + suffix
                    result[i] = format_table_row(cells)

    if warned_no_table:
        print(
            "  Warning: No '## Paint Equivalents' section found in the markdown file.",
            file=sys.stderr,
        )

    return result, rows_filled, rows_skipped


# ---------------------------------------------------------------------------
# Recipe parser (for DOCX generation)
# ---------------------------------------------------------------------------


def parse_recipe(lines: list) -> dict:
    """Parse recipe markdown into a structured dict for DOCX generation."""
    recipe = {
        "title": "",
        "source": "",
        "source_type": "",
        "paint_brands": "",
        "source_brand": "",
        "tags": "",
        "notes": [],
        "equivalents": [],
        "steps": [],
        "tips": [],
        "variations": [],
        "wider_application": [],
    }

    meta_map = {
        "**Source:**": "source",
        "**Source type:**": "source_type",
        "**Paint brands:**": "paint_brands",
        "**Source brand:**": "source_brand",
        "**Tags:**": "tags",
    }

    current_section = None

    for line in lines:
        s = line.rstrip("\n")

        # Title (# but not ##)
        if re.match(r"^#\s+", s) and not re.match(r"^##\s+", s):
            recipe["title"] = s[2:].strip()
            continue

        # Metadata fields
        matched_meta = False
        for prefix, key in meta_map.items():
            if s.startswith(prefix):
                recipe[key] = s[len(prefix) :].strip()
                matched_meta = True
                break
        if matched_meta:
            continue

        # Section headings
        if re.match(r"^##\s+", s):
            heading = s[3:].strip().lower()
            if "note" in heading:
                current_section = "notes"
            elif "paint equiv" in heading:
                current_section = "equivalents"
            elif "step" in heading:
                current_section = "steps"
            elif "tip" in heading:
                current_section = "tips"
            elif "variation" in heading:
                current_section = "variations"
            elif "wider" in heading or "application" in heading:
                current_section = "wider_application"
            elif "printable" in heading:
                current_section = None
            else:
                current_section = None
            continue

        # Accumulate content per section
        if current_section == "equivalents":
            cells = parse_table_row(s)
            if (
                cells
                and len(cells) >= 4
                and not is_separator_row(cells)
                and cells[0].lower() != "role"
            ):
                if len(cells) == 6:
                    recipe["equivalents"].append(
                        {
                            "role": cells[0],
                            "brand": cells[1],
                            "source_paint": cells[2],
                            "ttc": cells[3],
                            "citadel": cells[4],
                            "warpaints_fanatic": cells[5],
                        }
                    )
                elif len(cells) == 5:
                    recipe["equivalents"].append(
                        {
                            "role": cells[0],
                            "brand": "",
                            "source_paint": cells[1],
                            "ttc": cells[2],
                            "citadel": cells[3],
                            "warpaints_fanatic": cells[4],
                        }
                    )
                elif len(cells) == 4:
                    recipe["equivalents"].append(
                        {
                            "role": cells[0],
                            "brand": "",
                            "source_paint": cells[1],
                            "ttc": cells[2],
                            "citadel": "",
                            "warpaints_fanatic": cells[3],
                        }
                    )
        elif current_section and current_section != "equivalents":
            if s and not s.startswith("|") and not s.startswith("#"):
                recipe[current_section].append(s)

    return recipe


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------


def generate_docx(recipe: dict, output_path: Path, verbose: bool = False) -> None:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    COLOR_DARK = RGBColor(0x2D, 0x5A, 0x27)
    COLOR_TEXT = RGBColor(0x44, 0x44, 0x44)
    COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    HEX_DARK = "2D5A27"
    HEX_LIGHT = "EAF2E8"

    def set_cell_shading(cell, fill_hex: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)

    def add_run(para, text, bold=False, color=None, size_pt=None):
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        if size_pt:
            run.font.size = Pt(size_pt)
        return run

    def spacing(para, before=0, after=2):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)

    doc = Document()

    # Page: US Letter, 0.4" margins
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(0.4))

    doc.styles["Normal"].font.name = "Arial"  # type: ignore[union-attr]
    doc.styles["Normal"].font.size = Pt(9)  # type: ignore[union-attr]

    # Title
    p = doc.add_paragraph()
    spacing(p, before=0, after=4)
    add_run(
        p, recipe["title"] or "Untitled Recipe", bold=True, color=COLOR_DARK, size_pt=14
    )

    # Metadata line
    meta_parts = []
    if recipe["source"]:
        meta_parts.append(f"Source: {recipe['source']}")
    if recipe["source_type"]:
        meta_parts.append(f"Type: {recipe['source_type']}")
    if recipe["paint_brands"]:
        meta_parts.append(f"Brands: {recipe['paint_brands']}")
    if recipe["tags"]:
        meta_parts.append(f"Tags: {recipe['tags']}")
    if meta_parts:
        p = doc.add_paragraph()
        spacing(p, after=4)
        add_run(p, "  ·  ".join(meta_parts), color=COLOR_TEXT, size_pt=8)

    # Notes
    if recipe["notes"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Notes", bold=True, color=COLOR_DARK, size_pt=10)
        note_text = " ".join(
            ln.lstrip("- ").strip() for ln in recipe["notes"] if ln.strip()
        )
        p = doc.add_paragraph()
        spacing(p, after=4)
        add_run(p, note_text, color=COLOR_TEXT, size_pt=9)

    # Paint Equivalents table
    if recipe["equivalents"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Paint Equivalents", bold=True, color=COLOR_DARK, size_pt=10)

        headers = [
            "Role",
            "Source Paint",
            "Two Thin Coats",
            "Citadel",
            "Warpaints Fanatic",
        ]
        col_widths = [Inches(0.9), Inches(1.4), Inches(1.4), Inches(1.4), Inches(1.3)]

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header row
        hdr_row = table.rows[0]
        for idx, (hdr, width) in enumerate(zip(headers, col_widths)):
            cell = hdr_row.cells[idx]
            cell.width = width
            set_cell_shading(cell, HEX_DARK)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, hdr, bold=True, color=COLOR_WHITE, size_pt=8)

        # Data rows
        for row_idx, eq in enumerate(recipe["equivalents"]):
            row = table.add_row()
            fill = HEX_LIGHT if row_idx % 2 == 0 else "FFFFFF"
            values = [
                eq["role"],
                eq["source_paint"],
                eq["ttc"],
                eq["citadel"],
                eq["warpaints_fanatic"],
            ]
            for idx, (val, width) in enumerate(zip(values, col_widths)):
                cell = row.cells[idx]
                cell.width = width
                set_cell_shading(cell, fill)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                add_run(p, val, color=COLOR_TEXT, size_pt=8)

    # Steps
    if recipe["steps"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Steps", bold=True, color=COLOR_DARK, size_pt=10)

        step_num = 0
        for step_line in recipe["steps"]:
            s = step_line.strip()
            if not s:
                continue
            m = re.match(r"^\d+\.\s+\*\*(.+?)\*\*\s*[—\-]\s*(.+)$", s)
            if m:
                step_num += 1
                p = doc.add_paragraph()
                spacing(p, after=2)
                add_run(
                    p,
                    f"{step_num}. {m.group(1)} — ",
                    bold=True,
                    color=COLOR_DARK,
                    size_pt=9,
                )
                add_run(p, m.group(2), color=COLOR_TEXT, size_pt=9)
            else:
                p = doc.add_paragraph()
                spacing(p, after=2)
                add_run(p, s, color=COLOR_TEXT, size_pt=9)

    # Tips
    if recipe["tips"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Tips", bold=True, color=COLOR_DARK, size_pt=10)
        for tip in recipe["tips"]:
            t = tip.strip().lstrip("- ").strip()
            if t:
                p = doc.add_paragraph()
                spacing(p, after=1)
                add_run(p, f"• {t}", color=COLOR_TEXT, size_pt=9)

    # Variations & Ideas
    if recipe["variations"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Variations & Ideas", bold=True, color=COLOR_DARK, size_pt=10)
        for v in recipe["variations"]:
            t = v.strip().lstrip("- ").strip()
            if t:
                p = doc.add_paragraph()
                spacing(p, after=1)
                add_run(p, f"• {t}", color=COLOR_TEXT, size_pt=9)

    # Wider Application
    if recipe["wider_application"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Wider Application", bold=True, color=COLOR_DARK, size_pt=10)
        for w in recipe["wider_application"]:
            t = w.strip().lstrip("- ").strip()
            if t:
                p = doc.add_paragraph()
                spacing(p, after=1)
                add_run(p, f"• {t}", color=COLOR_TEXT, size_pt=9)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    if verbose:
        print(f"  DOCX saved: {output_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description=(
            "Fill paint equivalents in recipe markdown file(s) and generate DOCX."
        )
    )
    parser.add_argument(
        "recipe",
        nargs="?",
        help=(
            "Path to a single recipe .md file, or a folder containing recipe "
            ".md files (processes all .md files in the folder)"
        ),
    )
    parser.add_argument(
        "--recipes-dir",
        type=Path,
        help="Directory containing recipe markdown files (for batch mode)",
    )
    parser.add_argument(
        "--printables-dir",
        type=Path,
        help="Directory for output DOCX files (used with --recipes-dir)",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Re-fill all rows, overwriting any existing equivalent values",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Print detailed output (DOCX paths, row counts)",
    )
    args = parser.parse_args()

    script_dir = Path(__file__).parent
    paints = load_paints(script_dir)

    # Resolve batch target: positional folder arg takes precedence over --recipes-dir
    batch_dir: Path | None = None
    if args.recipe and Path(args.recipe).is_dir():
        batch_dir = Path(args.recipe)
    elif args.recipes_dir:
        batch_dir = args.recipes_dir

    # Batch mode: process all recipes in a directory
    if batch_dir:
        args.recipes_dir = batch_dir  # normalise for the block below

    if args.recipes_dir:
        recipes_dir = args.recipes_dir
        if not recipes_dir.exists():
            print(f"Error: Recipes directory not found: {recipes_dir}", file=sys.stderr)
            sys.exit(1)
        if not recipes_dir.is_dir():
            print(f"Error: Not a directory: {recipes_dir}", file=sys.stderr)
            sys.exit(1)

        printables_dir = (
            args.printables_dir
            if args.printables_dir
            else recipes_dir.parent / "Printables"
        )
        printables_dir.mkdir(parents=True, exist_ok=True)

        recipe_files = sorted(recipes_dir.glob("*.md"))

        if not recipe_files:
            print(f"No markdown files found in {recipes_dir}")
            sys.exit(0)

        force_note = " (force)" if args.force else ""
        print(
            f"Processing {len(recipe_files)} recipe(s){force_note} → {printables_dir}"
        )

        changed = 0
        unchanged = 0
        errors = 0

        for recipe_path in recipe_files:
            try:
                with open(recipe_path, "r", encoding="utf-8") as f:
                    lines = f.readlines()

                updated_lines, rows_filled, rows_skipped = fill_equivalents(
                    lines, paints, args.force
                )

                # Check if markdown actually changed
                md_changed = lines != updated_lines

                with open(recipe_path, "w", encoding="utf-8") as f:
                    f.writelines(updated_lines)

                docx_path = printables_dir / (recipe_path.stem + ".docx")
                recipe_data = parse_recipe(updated_lines)
                generate_docx(recipe_data, docx_path, verbose=args.verbose)

                if md_changed:
                    if args.verbose:
                        print(
                            f"  {recipe_path.name}: filled {rows_filled}, skipped {rows_skipped}"
                        )
                    else:
                        print(f"  {recipe_path.name}")
                    changed += 1
                else:
                    unchanged += 1

            except Exception as e:
                print(f"  [ERROR] {recipe_path.name}: {e}")
                errors += 1

        parts = []
        if changed:
            parts.append(f"{changed} changed")
        if unchanged:
            parts.append(f"{unchanged} unchanged")
        if errors:
            parts.append(f"{errors} errors")
        print(f"Done. {', '.join(parts)}.")
        return

    # Single file mode
    if not args.recipe:
        parser.error(
            "Either provide a recipe file or folder, or use --recipes-dir for batch mode"
        )

    recipe_path = Path(args.recipe)
    if not recipe_path.exists():
        print(f"Error: Recipe file not found: {recipe_path}", file=sys.stderr)
        sys.exit(1)
    if recipe_path.suffix.lower() != ".md":
        print(
            f"Warning: Expected a .md file, got: {recipe_path.suffix}",
            file=sys.stderr,
        )

    print(f"Processing: {recipe_path}")

    with open(recipe_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    updated_lines, rows_filled, rows_skipped = fill_equivalents(
        lines, paints, args.force
    )

    # Write the updated markdown back in-place
    with open(recipe_path, "w", encoding="utf-8") as f:
        f.writelines(updated_lines)

    if args.verbose:
        print(f"  Rows filled: {rows_filled}, rows skipped: {rows_skipped}")

    # Derive DOCX path: sibling Printables/ directory relative to the recipe
    printables_dir = recipe_path.parent.parent / "Printables"
    docx_path = printables_dir / (recipe_path.stem + ".docx")

    if args.verbose:
        print(f"Generating DOCX: {docx_path}")
    try:
        recipe_data = parse_recipe(updated_lines)
        generate_docx(recipe_data, docx_path, verbose=args.verbose)
    except ImportError:
        print(
            "Error: python-docx is not installed. Install with: pip install python-docx",
            file=sys.stderr,
        )
        sys.exit(1)
    except Exception as e:
        print(f"Error generating DOCX: {e}", file=sys.stderr)
        sys.exit(1)

    print("Done.")


if __name__ == "__main__":
    main()
