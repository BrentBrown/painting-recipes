#!/usr/bin/env python3
"""
convert_all_recipes.py

Batch convert all recipe markdown files in the Recipes folder to DOCX files
in the Printables folder. Idempotent by default - only converts files that
don't already have a corresponding DOCX.

Usage:
    python convert_all_recipes.py
    python convert_all_recipes.py --force
    python convert_all_recipes.py --recipes-dir ./Recipes --printables-dir ./Printables
"""

import argparse
import json
import re
import sys
from pathlib import Path

NO_EQ = "No equivalent"

SUFFIX_PATTERN = re.compile(
    r"\s+(Gloss|Contrast|Shade|Wash|Matte|Matt|Layer|Base|Dry|Technical)$",
    re.IGNORECASE,
)


def strip_suffixes(name: str) -> str:
    stripped = name.strip()
    while True:
        new = SUFFIX_PATTERN.sub("", stripped)
        if new == stripped:
            break
        stripped = new
    return stripped


def load_paints(script_dir: Path) -> dict:
    paints_path = script_dir / "paints.json"
    if not paints_path.exists():
        print(
            f"Error: paints.json not found at {paints_path}\n"
            f"Expected it alongside this script in: {script_dir}",
            file=sys.stderr,
        )
        sys.exit(1)
    try:
        with open(paints_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        print(f"Error: paints.json is malformed: {e}", file=sys.stderr)
        sys.exit(1)


def build_wf_reverse_lookup(paints: dict) -> dict:
    wf_table = paints.get("Warpaints Fanatic", {})
    reverse: dict[str, list[str]] = {}
    for wf_name, entry in wf_table.items():
        cit = entry.get("citadel")
        if cit and cit != NO_EQ:
            reverse.setdefault(cit, []).append(wf_name)
    return {cit: " / ".join(names) for cit, names in reverse.items()}


def lookup_equivalents(
    brand: str, source_paint: str, paints: dict, wf_reverse: dict
) -> tuple:
    stripped = strip_suffixes(source_paint)

    brand_table = paints.get(brand)

    # --- Special handling: Check Citadel Contrast when brand is Citadel ---
    if brand == "Citadel":
        contrast_table = paints.get("Citadel Contrast")
        if contrast_table:
            contrast_entry = contrast_table.get(stripped) or contrast_table.get(source_paint.strip())
            if contrast_entry:
                speedpaint = contrast_entry.get("speedpaint", NO_EQ)
                return NO_EQ, NO_EQ, speedpaint

    if brand_table is None:
        return NO_EQ, NO_EQ, NO_EQ

    entry = brand_table.get(stripped) or brand_table.get(source_paint.strip())

    if entry is None:
        return NO_EQ, NO_EQ, NO_EQ

    # --- Special handling for Citadel Contrast ---
    if brand == "Citadel Contrast":
        speedpaint = entry.get("speedpaint", NO_EQ)
        return NO_EQ, NO_EQ, speedpaint

    # --- Special handling for Army Painter Speedpaint ---
    if brand == "Army Painter Speedpaint":
        contrast = entry.get("citadel_contrast", NO_EQ)
        return NO_EQ, NO_EQ, contrast

    if brand == "Two Thin Coats":
        wave = entry.get("wave", "")
        ttc_col = f"{stripped} ({wave})" if wave else stripped
    else:
        ttc_name = entry.get("ttc") or NO_EQ
        if ttc_name == NO_EQ:
            ttc_col = NO_EQ
        else:
            wave = entry.get("wave", "")
            ttc_col = f"{ttc_name} ({wave})" if wave else ttc_name

    if brand == "Citadel":
        citadel_col = source_paint.strip()
    else:
        citadel_col = entry.get("citadel") or NO_EQ

    if brand == "Warpaints Fanatic":
        wf_col = source_paint.strip()
    else:
        wf_col = wf_reverse.get(citadel_col, NO_EQ)

    return ttc_col, citadel_col, wf_col


def parse_table_row(line: str) -> list | None:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return None
    cells = [c.strip() for c in stripped.split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def is_separator_row(cells: list) -> bool:
    return bool(cells) and all(re.match(r"^:?-+:?$", c) for c in cells)


def format_table_row(cells: list) -> str:
    return "| " + " | ".join(cells) + " |\n"


# Speedpaint names for detection
SPEEDPAINT_NAMES = None


def get_speedpaint_names(paints: dict) -> set:
    """Get set of all Speedpaint paint names for detection."""
    global SPEEDPAINT_NAMES
    if SPEEDPAINT_NAMES is None:
        sp_table = paints.get("Army Painter Speedpaint", {})
        SPEEDPAINT_NAMES = set(sp_table.keys())
    return SPEEDPAINT_NAMES


def is_speedpaint(paint_name: str, speedpaint_names: set) -> bool:
    """Check if a paint name is a Speedpaint."""
    if not paint_name or paint_name == NO_EQ:
        return False
    return paint_name in speedpaint_names


def fill_equivalents(lines: list, paints: dict, force: bool) -> tuple:
    result = []
    in_equiv_section = False
    header_seen = False
    rows_filled = 0
    rows_skipped = 0
    warned_no_table = True
    wf_reverse = build_wf_reverse_lookup(paints)
    table_columns = 0

    for i, line in enumerate(lines):
        if re.match(r"^##\s+Paint Equivalents", line.strip()):
            in_equiv_section = True
            header_seen = False
            warned_no_table = False
            table_columns = 0
            result.append(line)
            continue

        if in_equiv_section and re.match(r"^##\s+", line.strip()):
            in_equiv_section = False

        if not in_equiv_section:
            result.append(line)
            continue

        cells = parse_table_row(line)

        if cells is None:
            result.append(line)
            continue

        if is_separator_row(cells):
            result.append(line)
            continue

        if not header_seen:
            if len(cells) >= 5 and cells[0].lower() == "role":
                header_seen = True
                table_columns = len(cells)
            result.append(line)
            continue

        if len(cells) != table_columns:
            result.append(line)
            rows_skipped += 1
            continue

        if table_columns == 6:
            role, brand, source_paint, ttc_val, citadel_val, wf_val = cells
        elif table_columns == 5:
            role, source_paint, ttc_val, citadel_val, wf_val = cells
            brand = ""
        else:
            result.append(line)
            rows_skipped += 1
            continue

        if not source_paint or not brand:
            result.append(line)
            continue

        all_filled = bool(ttc_val and citadel_val and wf_val)

        if all_filled and not force:
            rows_skipped += 1
            result.append(line)
            continue

        ttc_new, citadel_new, wf_new = lookup_equivalents(
            brand, source_paint, paints, wf_reverse
        )

        if table_columns == 6:
            new_cells = [role, brand, source_paint, ttc_new, citadel_new, wf_new]
        else:
            new_cells = [role, source_paint, ttc_new, citadel_new, wf_new]

        result.append(format_table_row(new_cells))
        rows_filled += 1

    # Second: scan result to determine header based on wf column values
    speedpaint_names = get_speedpaint_names(paints)
    has_fanatic = False
    has_speedpaint = False

    for line in result:
        cells = parse_table_row(line)
        if cells and len(cells) >= 6 and cells[0].lower() != "role":
            wf_val = cells[5]
            if is_speedpaint(wf_val, speedpaint_names):
                has_speedpaint = True
            elif wf_val and wf_val != NO_EQ:
                has_fanatic = True

    # Determine header text
    if has_speedpaint and has_fanatic:
        new_header = "Warpaints Fanatic / Speedpaint 2.0"
    elif has_speedpaint:
        new_header = "Speedpaint 2.0"
    else:
        new_header = None  # Keep original header

    # Update header only if we have Speedpaint (otherwise keep original)
    if new_header and table_columns == 6 and header_seen:
        for i, line in enumerate(result):
            cells = parse_table_row(line)
            if cells and len(cells) >= 6 and cells[0].lower() == "role":
                cells[5] = new_header
                result[i] = format_table_row(cells)
                break

    return result, rows_filled, rows_skipped


def parse_recipe(lines: list) -> dict:
    recipe = {
        "title": "",
        "source": "",
        "source_type": "",
        "paint_brands": "",
        "source_brand": "",
        "tags": "",
        "notes": [],
        "equivalents": [],
        "steps": [],
        "tips": [],
        "variations": [],
        "wider_application": [],
    }

    meta_map = {
        "**Source:**": "source",
        "**Source type:**": "source_type",
        "**Paint brands:**": "paint_brands",
        "**Source brand:**": "source_brand",
        "**Tags:**": "tags",
    }

    current_section = None

    for line in lines:
        s = line.rstrip("\n")

        if re.match(r"^#\s+", s) and not re.match(r"^##\s+", s):
            recipe["title"] = s[2:].strip()
            continue

        matched_meta = False
        for prefix, key in meta_map.items():
            if s.startswith(prefix):
                recipe[key] = s[len(prefix) :].strip()
                matched_meta = True
                break
        if matched_meta:
            continue

        if re.match(r"^##\s+", s):
            heading = s[3:].strip().lower()
            if "note" in heading:
                current_section = "notes"
            elif "paint equiv" in heading:
                current_section = "equivalents"
            elif "step" in heading:
                current_section = "steps"
            elif "tip" in heading:
                current_section = "tips"
            elif "variation" in heading:
                current_section = "variations"
            elif "wider" in heading or "application" in heading:
                current_section = "wider_application"
            elif "printable" in heading:
                current_section = None
            else:
                current_section = None
            continue

        if current_section == "equivalents":
            cells = parse_table_row(s)
            if cells and len(cells) >= 5 and not is_separator_row(cells) and cells[0].lower() != "role":
                if len(cells) == 6:
                    recipe["equivalents"].append(
                        {
                            "role": cells[0],
                            "brand": cells[1],
                            "source_paint": cells[2],
                            "ttc": cells[3],
                            "citadel": cells[4],
                            "warpaints_fanatic": cells[5],
                        }
                    )
                elif len(cells) == 5:
                    recipe["equivalents"].append(
                        {
                            "role": cells[0],
                            "brand": "",
                            "source_paint": cells[1],
                            "ttc": cells[2],
                            "citadel": cells[3],
                            "warpaints_fanatic": cells[4],
                        }
                    )
        elif current_section and current_section != "equivalents":
            if s and not s.startswith("|") and not s.startswith("#"):
                recipe[current_section].append(s)

    return recipe


def generate_docx(recipe: dict, output_path: Path) -> None:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    COLOR_DARK = RGBColor(0x2D, 0x5A, 0x27)
    COLOR_TEXT = RGBColor(0x44, 0x44, 0x44)
    COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    HEX_DARK = "2D5A27"
    HEX_LIGHT = "EAF2E8"

    def set_cell_shading(cell, fill_hex: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)

    def add_run(para, text, bold=False, color=None, size_pt=None):
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        if size_pt:
            run.font.size = Pt(size_pt)
        return run

    def spacing(para, before=0, after=2):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)

    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(0.4))

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)

    p = doc.add_paragraph()
    spacing(p, before=0, after=4)
    add_run(
        p, recipe["title"] or "Untitled Recipe", bold=True, color=COLOR_DARK, size_pt=14
    )

    meta_parts = []
    if recipe["source"]:
        meta_parts.append(f"Source: {recipe['source']}")
    if recipe["source_type"]:
        meta_parts.append(f"Type: {recipe['source_type']}")
    if recipe["paint_brands"]:
        meta_parts.append(f"Brands: {recipe['paint_brands']}")
    if recipe["tags"]:
        meta_parts.append(f"Tags: {recipe['tags']}")
    if meta_parts:
        p = doc.add_paragraph()
        spacing(p, after=4)
        add_run(p, "  ·  ".join(meta_parts), color=COLOR_TEXT, size_pt=8)

    if recipe["notes"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Notes", bold=True, color=COLOR_DARK, size_pt=10)
        note_text = " ".join(
            ln.lstrip("- ").strip() for ln in recipe["notes"] if ln.strip()
        )
        p = doc.add_paragraph()
        spacing(p, after=4)
        add_run(p, note_text, color=COLOR_TEXT, size_pt=9)

    if recipe["equivalents"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Paint Equivalents", bold=True, color=COLOR_DARK, size_pt=10)

        headers = [
            "Role",
            "Source Paint",
            "Two Thin Coats",
            "Citadel",
            "Warpaints Fanatic",
        ]
        col_widths = [Inches(0.9), Inches(1.4), Inches(1.4), Inches(1.4), Inches(1.3)]

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        hdr_row = table.rows[0]
        for idx, (hdr, width) in enumerate(zip(headers, col_widths)):
            cell = hdr_row.cells[idx]
            cell.width = width
            set_cell_shading(cell, HEX_DARK)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, hdr, bold=True, color=COLOR_WHITE, size_pt=8)

        for row_idx, eq in enumerate(recipe["equivalents"]):
            row = table.add_row()
            fill = HEX_LIGHT if row_idx % 2 == 0 else "FFFFFF"
            values = [
                eq["role"],
                eq["source_paint"],
                eq["ttc"],
                eq["citadel"],
                eq["warpaints_fanatic"],
            ]
            for idx, (val, width) in enumerate(zip(values, col_widths)):
                cell = row.cells[idx]
                cell.width = width
                set_cell_shading(cell, fill)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                add_run(p, val, color=COLOR_TEXT, size_pt=8)

    if recipe["steps"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Steps", bold=True, color=COLOR_DARK, size_pt=10)

        step_num = 0
        for step_line in recipe["steps"]:
            s = step_line.strip()
            if not s:
                continue
            m = re.match(r"^\d+\.\s+\*\*(.+?)\*\*\s*[—\-]\s*(.+)$", s)
            if m:
                step_num += 1
                p = doc.add_paragraph()
                spacing(p, after=2)
                add_run(
                    p,
                    f"{step_num}. {m.group(1)} — ",
                    bold=True,
                    color=COLOR_DARK,
                    size_pt=9,
                )
                add_run(p, m.group(2), color=COLOR_TEXT, size_pt=9)
            else:
                p = doc.add_paragraph()
                spacing(p, after=2)
                add_run(p, s, color=COLOR_TEXT, size_pt=9)

    if recipe["tips"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Tips", bold=True, color=COLOR_DARK, size_pt=10)
        for tip in recipe["tips"]:
            t = tip.strip().lstrip("- ").strip()
            if t:
                p = doc.add_paragraph()
                spacing(p, after=1)
                add_run(p, f"• {t}", color=COLOR_TEXT, size_pt=9)

    if recipe["variations"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Variations & Ideas", bold=True, color=COLOR_DARK, size_pt=10)
        for v in recipe["variations"]:
            t = v.strip().lstrip("- ").strip()
            if t:
                p = doc.add_paragraph()
                spacing(p, after=1)
                add_run(p, f"• {t}", color=COLOR_TEXT, size_pt=9)

    if recipe["wider_application"]:
        p = doc.add_paragraph()
        spacing(p, before=4, after=2)
        add_run(p, "Wider Application", bold=True, color=COLOR_DARK, size_pt=10)
        for w in recipe["wider_application"]:
            t = w.strip().lstrip("- ").strip()
            if t:
                p = doc.add_paragraph()
                spacing(p, after=1)
                add_run(p, f"• {t}", color=COLOR_TEXT, size_pt=9)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def process_recipe(recipe_path: Path, printables_dir: Path, paints: dict, force: bool) -> dict:
    result = {
        "path": recipe_path,
        "action": "skipped",
        "reason": "",
        "error": None,
    }

    docx_path = printables_dir / (recipe_path.stem + ".docx")

    if docx_path.exists() and not force:
        result["reason"] = "DOCX already exists"
        return result

    try:
        with open(recipe_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        updated_lines, rows_filled, rows_skipped = fill_equivalents(
            lines, paints, force
        )

        with open(recipe_path, "w", encoding="utf-8") as f:
            f.writelines(updated_lines)

        recipe_data = parse_recipe(updated_lines)
        generate_docx(recipe_data, docx_path)

        result["action"] = "converted" if not docx_path.exists() or force else "updated"
        result["reason"] = f"Rows filled: {rows_filled}, skipped: {rows_skipped}"

    except Exception as e:
        result["error"] = str(e)

    return result


def main():
    parser = argparse.ArgumentParser(
        description="Batch convert recipe markdown files to DOCX"
    )
    parser.add_argument(
        "--recipes-dir",
        type=Path,
        default=Path("recipes"),
        help="Directory containing recipe markdown files",
    )
    parser.add_argument(
        "--printables-dir",
        type=Path,
        default=Path("printables"),
        help="Directory for output DOCX files",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Re-convert even if DOCX already exists",
    )
    args = parser.parse_args()

    project_root = Path(__file__).parent.parent

    if args.recipes_dir.is_absolute():
        recipes_dir = args.recipes_dir
    else:
        recipes_dir = project_root / args.recipes_dir

    if args.printables_dir.is_absolute():
        printables_dir = args.printables_dir
    else:
        printables_dir = project_root / args.printables_dir

    if not recipes_dir.exists():
        print(f"Error: Recipes directory not found: {recipes_dir}", file=sys.stderr)
        sys.exit(1)

    if not recipes_dir.is_dir():
        print(f"Error: Not a directory: {recipes_dir}", file=sys.stderr)
        sys.exit(1)

    script_dir = Path(__file__).parent
    paints = load_paints(script_dir)

    recipe_files = sorted(recipes_dir.glob("*.md"))

    if not recipe_files:
        print(f"No markdown files found in {recipes_dir}")
        sys.exit(0)

    printables_dir.mkdir(parents=True, exist_ok=True)

    print(f"Recipes dir: {recipes_dir}")
    print(f"Printables dir: {printables_dir}")
    print(f"Found {len(recipe_files)} recipe(s)")
    print(f"Force mode: {args.force}")
    print("-" * 50)

    converted = 0
    skipped = 0
    errors = 0

    for recipe_path in recipe_files:
        try:
            result = process_recipe(recipe_path, printables_dir, paints, args.force)

            if result["error"]:
                print(f"  [ERROR] {recipe_path.name}: {result['error']}")
                errors += 1
            elif result["action"] == "skipped":
                print(f"  [SKIP]  {recipe_path.name}: {result['reason']}")
                skipped += 1
            else:
                print(f"  [OK]    {recipe_path.name}: {result['reason']}")
                converted += 1

        except Exception as e:
            print(f"  [ERROR] {recipe_path.name}: {e}")
            errors += 1

    print("-" * 50)
    print(f"Summary: {converted} converted, {skipped} skipped, {errors} errors")


if __name__ == "__main__":
    main()
